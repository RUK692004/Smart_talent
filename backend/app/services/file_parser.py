import os
from typing import List
import re
import pdfplumber
import pytesseract
from docx import Document
from PIL import Image
from app.preprocess.image_preprocessor import preprocess_image_for_ocr

# ---------- Configuration ----------
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

SUPPORTED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}


# ---------- Helpers ----------
def join_non_empty(lines: List[str]) -> str:
    return "\n".join(line.strip() for line in lines if line and line.strip()).strip()


# ---------- Main Entry ----------
def extract_text(file_path: str) -> str:
    """
    Detect file type and extract text using the appropriate method.
    """
    try:
        _, ext = os.path.splitext(file_path.lower())
        print(f"FILE_PARSER: detected extension = {ext}")

        if ext == ".pdf":
            print("FILE_PARSER: using PDF text extraction")
            return extract_text_from_pdf(file_path)

        elif ext == ".docx":
            print("FILE_PARSER: using DOCX direct extraction")
            return extract_text_from_docx(file_path)

        elif ext in SUPPORTED_IMAGE_EXTENSIONS:
            print("FILE_PARSER: using image OCR extraction")
            return extract_text_from_image(file_path)

        else:
            raise ValueError(f"Unsupported file format: {ext}")

    except Exception as e:
        raise Exception(f"Error processing file '{file_path}': {str(e)}")


# ---------- PDF Extraction ----------
def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from text-based PDF files.
    """
    try:
        extracted_lines = []

        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    extracted_lines.append(text)
                else:
                    print(f"FILE_PARSER: no text found on PDF page {page_number}")

        final_text = join_non_empty(extracted_lines)

        if not final_text:
            raise Exception("No readable text found in PDF.")

        return final_text

    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


# ---------- DOCX Extraction ----------
def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text directly from DOCX files.
    """
    try:
        doc = Document(file_path)
        extracted_lines = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_lines.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)

                if row_data:
                    extracted_lines.append(" | ".join(row_data))

        final_text = join_non_empty(extracted_lines)

        if not final_text:
            raise Exception("No readable text found in DOCX.")

        return final_text

    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

# ---------- PSM6 / PSM11 ----------
def choose_best_ocr(text1: str, text2: str) -> str:
    """
    Choose the better OCR result based on simple heuristics.
    """

    def score(text: str) -> int:
        score = 0

        # Reward presence of email
        if re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.", text):
            score += 5

        # Reward phone-like numbers
        if re.search(r"\d{10}", text):
            score += 3

        # Reward known section headings
        keywords = [
            "experience", "education", "skills",
            "projects", "profile", "summary"
        ]

        for kw in keywords:
            if kw in text.lower():
                score += 2

        # Penalize too much garbage characters
        garbage = len(re.findall(r"[^a-zA-Z0-9\s.,\-@]", text))
        score -= garbage // 50

        return score

    score1 = score(text1)
    score2 = score(text2)

    print(f"OCR SCORE psm6: {score1}")
    print(f"OCR SCORE psm11: {score2}")

    return text1 if score1 >= score2 else text2 

# ---------- Image Extraction ----------
def extract_text_from_image(file_path: str) -> str:
    """
    Run OCR with multiple PSM modes and choose the best result.
    """
    try:
        print("FILE_PARSER: preprocessing image")
        processed_image = preprocess_image_for_ocr(file_path)

        print("FILE_PARSER: running OCR (psm 6)")
        text_psm6 = pytesseract.image_to_string(
            processed_image,
            config='--oem 3 --psm 6'
        )

        print("FILE_PARSER: running OCR (psm 11)")
        text_psm11 = pytesseract.image_to_string(
            processed_image,
            config='--oem 3 --psm 11'
        )

        # Choose better result
        best_text = choose_best_ocr(text_psm6, text_psm11)

        final_text = best_text.strip()

        if not final_text:
            raise Exception("No readable text found in image.")

        return final_text

    except Exception as e:
        raise Exception(f"OCR failed: {str(e)}")